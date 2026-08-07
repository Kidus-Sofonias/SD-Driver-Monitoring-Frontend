"""Generate a .docx presentation script for the Drive Pulse judges presentation."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Calibri'
    heading_style.font.color.rgb = RGBColor(0x0a, 0x6c, 0xb5)

# ── Cover Page ──────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Drive Pulse: AI Driver Monitor')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x0a, 0x6c, 0xb5)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Presentation Script — Judges Presentation')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Prepared for: Judges Panel\nDate: July 2026\nTeam: VisionZero — Addis Ababa, Ethiopia')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ── Helper ──────────────────────────────────────────────────────────────────
def add_slide_header(number, title_text, slide_title=""):
    """Add a slide header with slide number and title."""
    h = doc.add_heading(f'Slide {number}: {title_text}', level=1)
    if slide_title:
        p = doc.add_paragraph()
        run = p.add_run(f'Slide content title: "{slide_title}"')
        run.italic = True
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.size = Pt(10)
    doc.add_paragraph()

def add_script(script_text):
    """Add a presentation script paragraph."""
    p = doc.add_paragraph()
    run = p.add_run('🗣️ Script:')
    run.bold = True
    run.font.color.rgb = RGBColor(0x0a, 0x6c, 0xb5)
    p2 = doc.add_paragraph(script_text)
    p2.style.font.size = Pt(11)
    p2.paragraph_format.space_after = Pt(6)

def add_tip(tip_text):
    """Add a delivery tip."""
    p = doc.add_paragraph()
    run = p.add_run('💡 Delivery Tip: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xb5, 0x6c, 0x0a)
    run = p.add_run(tip_text)
    run.font.color.rgb = RGBColor(0x88, 0x66, 0x22)
    run.font.size = Pt(10)

def add_visual(desc):
    """Add visual description."""
    p = doc.add_paragraph()
    run = p.add_run('📺 Visual: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x6c, 0x0a, 0xb5)
    run = p.add_run(desc)
    run.font.size = Pt(10)

def add_separator():
    p = doc.add_paragraph()
    run = p.add_run('─' * 60)
    run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    run.font.size = Pt(8)

# ── Slide 1: Title / Introduction ──────────────────────────────────────────
add_slide_header(1, 'Title & Introduction', 'Drive Pulse: AI Driver Monitor')

add_visual('Dark slide with animated glow particles. Title "Drive Pulse" large centered, with tagline "AI Driver Monitor" underneath. Subtitle "Driver Intelligence" fades in. Team "VisionZero — Addis Ababa, Ethiopia" at bottom.')

add_script(
    'Good morning/afternoon, esteemed judges and panel members. My name is [Your Name], and I am honored to present Drive Pulse — '
    'an AI-powered driver monitoring system designed to tackle one of the most pressing yet overlooked crises in Ethiopia today: '
    'road safety. Over the next few minutes, I will walk you through the problem, our solution, what we have already built, and '
    'why this matters — not just for Ethiopia, but as a model for low-income countries worldwide.'
)

add_tip('Speak with confidence and deliberate pacing. This slide sets the tone — make eye contact with the judges as you deliver the opening sentence. Pause briefly after "road safety".')

add_separator()

# ── Slide 2: The Problem — Road Safety Crisis ──────────────────────────────
add_slide_header(2, 'The Road Safety Crisis', '1.16 million deaths — one every 27 seconds')

add_visual('Large number "1.16M" displayed prominently. A counter ticking every 27 seconds. World map highlighting Ethiopia in red. Pull quotes from WHO data.')

add_script(
    'Let me start with a staggering fact: every year, 1.16 million people die on the world\'s roads. That is one death every 27 seconds. '
    'To put that in perspective — by the time I finish this sentence, someone, somewhere, has lost their life. '
    'Road traffic injuries are the leading cause of death for young people aged 5 to 29. '
    'And 90% of these deaths happen in low- and middle-income countries — precisely the places least equipped to handle them.'
)

add_tip('Let the "one death every 27 seconds" line land. Pause for 2 seconds after saying it. The gravity of this statistic is your hook.')

add_separator()

# ── Slide 3: Ethiopia Specifically ─────────────────────────────────────────
add_slide_header(3, 'Ethiopia — Ground Zero', 'Fatality rate of 37 per 100,000 population')

add_visual('Bar chart comparing regional fatality rates per 100k: Ethiopia 37, Africa average 27, SE Asia 20, Global 18.2, Europe 9.3, Americas 16. Ethiopia\'s bar is highlighted in red.')

add_script(
    'Now, let\'s zoom in on Ethiopia. Our country has a road fatality rate of 37 per 100,000 population. That is more than double '
    'the global average of 18.2. To give you another metric — we have approximately 4,922 fatalities per 100,000 registered motor vehicles. '
    'Young adults aged 15 to 29 — the most economically productive segment of our population — are at the highest risk. '
    'Pedestrians account for roughly 30% of fatalities, and motorcyclists another 16%. '
    'This is not just a health crisis — it is an economic one, draining hundreds of millions of dollars from our economy each year.'
)

add_tip('Point to Ethiopia on the chart as you mention it. The contrast with Europe (9.3) is powerful — emphasize it.')

add_separator()

# ── Slide 4: Common Causes ─────────────────────────────────────────────────
add_slide_header(4, 'What\'s Causing the Crashes?', 'Speeding: 38% | Phone distraction: 22% | Alcohol: 18%')

add_visual('Horizontal bar chart showing: Speeding/reckless driving 38%, Distracted driving (phone) 22%, Drunk/impaired driving 18%, Poor infrastructure 14%, Vehicle defects 8%. Phone distraction section is highlighted.')

add_script(
    'When we analyze the root causes, a clear picture emerges. Speeding and reckless driving account for 38% of fatalities. '
    'But what I want to draw your attention to is the second largest cause: distracted driving — specifically phone use — at 22%. '
    'And here is why that number is so concerning: research shows that dialing a phone while driving increases crash risk by 2.7 times. '
    'Reading a text? 4.8 times. Sending a text? 6.1 times. And reaching for a phone — just reaching for it — increases crash risk by more than 12 times. '
    'The simple act of looking at your phone for two seconds doubles your risk of a crash.'
)

add_tip('Use your hand to gesture "reaching" when you mention the 12x risk. This visual cue helps the judges feel the danger.')

add_separator()

# ── Slide 5: The Technology Gap ────────────────────────────────────────────
add_slide_header(5, 'Why Existing Solutions Fail', '$100–500 per vehicle + monthly fees + professional installation')

add_visual('Three panels showing: (1) Traditional telematics box with high price tag, (2) Insurance black box with privacy lock icon, (3) Smartphone with checkmark. Below: "88% smartphone penetration in Africa by 2030" with rising graph.')

add_script(
    'You might ask: why doesn\'t existing technology solve this? The answer is cost and accessibility. '
    'Traditional telematics systems cost between 100 and 500 dollars per vehicle, plus monthly subscription fees, plus professional installation. '
    'Insurance black boxes are tied to insurers, raise privacy concerns, and provide no feedback to the driver. '
    'None of these solutions are designed for low-income markets. '
    'But here is the opportunity: smartphone penetration in Africa is projected to reach 88% by 2030. '
    'Every single one of those phones already has a GPS, an accelerometer, and a gyroscope. '
    'The hardware is already in people\'s pockets — we just needed the software to make it work.'
)

add_tip('This is your "aha moment" slide. Emphasize "the hardware is already in people\'s pockets" with a slight pause before and after.')

add_separator()

# ── Slide 6: Our Solution — Drive Pulse ─────────────────────────────────────
add_slide_header(6, 'Introducing Drive Pulse', 'Zero extra hardware — just a phone')

add_visual('Phone mockup showing the Drive Pulse app in "Live Trip" mode. Animated sensor data flowing. Three language badges: EN, AM, OM. Split screen with driver view and fleet admin view.')

add_script(
    'This is where Drive Pulse comes in. We have built a system that turns a standard smartphone into a professional-grade driving safety tool. '
    'Zero extra hardware. No installation costs. No monthly fees hidden in fine print. Just a phone and our app. '
    'Here is how it works: the driver opens the app, starts a trip, and the phone\'s built-in sensors — GPS, accelerometer, and gyroscope — '
    'begin collecting data. That data is uploaded to our backend, where both machine learning models and rule-based scoring algorithms '
    'analyze the driving behavior. Within moments, the driver receives a safety score, a risk level, a confidence indicator, and — '
    'most importantly — plain-language reasons explaining what they did well and what they can improve. '
    'And it works in English, Amharic, and Oromo, right out of the box.'
)

add_tip('If you have a phone with the app, hold it up here as a prop. Show them the live trip screen.')

add_separator()

# ── Slide 7: How It Works (Process) ────────────────────────────────────────
add_slide_header(7, 'How It Works — Step by Step', 'Start → Collect → Upload → Analyze → Score')

add_visual('Five-step horizontal flow diagram: (1) Start Trip, (2) Collect Sensor Data, (3) Upload Samples, (4) AI Analysis, (5) Safety Score & Insights. Each step has an icon.')

add_script(
    'Let me walk you through the process step by step. Step one: a driver opens the app and taps to start a new trip session. '
    'Step two: the phone\'s sensors begin sampling — GPS for location and speed, accelerometer for acceleration and braking force, '
    'and gyroscope for steering and turning behavior. Step three: those sensor readings are buffered and uploaded to our cloud backend '
    'with smart status-aware handling that works even in areas with spotty connectivity. '
    'Step four: the backend runs it through our analysis pipeline — a combination of trained machine learning models and deterministic '
    'rule-based scoring. And step five: the driver receives a comprehensive safety score, risk classification, confidence band, '
    'and actionable reasons for the rating. Every trip becomes a learning opportunity.'
)

add_tip('Move your hand across the screen as you describe each step. The flow should feel natural and inevitable.')

add_separator()

# ── Slide 8: Technical Architecture ─────────────────────────────────────────
add_slide_header(8, 'What We\'ve Built — Behind the Scenes', 'Full-stack: React Native → FastAPI → ML Pipeline → PostgreSQL')

add_visual('Architecture diagram: Mobile App (React Native/Expo) → API Layer (FastAPI on Render) → ML Pipeline (scikit-learn + ONNX) → Database (Supabase PostgreSQL). Arrows showing data flow.')

add_script(
    'Behind the scenes, this is a full-stack production system. The mobile app is built with React Native and Expo, deployed as an Android APK '
    'available for direct download from our website. The backend is a FastAPI application running on Render, with a PostgreSQL database '
    'powered by Supabase. Our machine learning pipeline — built with scikit-learn and supported by ONNX for inference — handles everything '
    'from synthetic data generation through model training, evaluation, promotion, and automatic retraining. '
    'We detect six types of driving events, including hard braking, rapid acceleration, sharp turns, and unstable motion. '
    'The system is designed to work even without a trained model — our rule-based scoring rules provide a fallback, '
    'so the product is never dependent on having ML available.'
)

add_tip('Don\'t get too technical here. Judges want to know you\'ve built something real — they don\'t need every detail. Keep this slide to about 45 seconds.')

add_separator()

# ── Slide 9: Pilot Results ─────────────────────────────────────────────────
add_slide_header(9, 'Pilot Results — Real Data', '50+ trips | 144,000+ samples | 6 drivers | 6 event types')

add_visual('Dashboard-style metrics panel showing: "50+ Trips Captured", "144K+ Sensor Samples", "6 Active Drivers", "3 Languages Supported", "6 Event Types Detected". Below: before/after behavior improvement chart.')

add_script(
    'We have already put this system to the test. Our pilot has captured over 50 real trips on Ethiopian roads, collecting more than '
    '144,000 individual sensor samples across 6 active driver accounts. The system successfully detects 6 distinct types of driving events. '
    'Our ML pipeline is fully operational — synthetic data generation, model training, evaluation, promotion to production, and automatic '
    'retraining all happen without manual intervention. The entire infrastructure is live on Render with Supabase PostgreSQL, '
    'and the Android APK is distributed directly from our website. This is not a prototype — this is a deployed, functioning system.'
)

add_tip('When you say "50+ trips" and "144,000 samples", let the numbers stand on their own. Don\'t downplay them — these are real achievements for a pilot.')

add_separator()

# ── Slide 10: What Drivers See ──────────────────────────────────────────────
add_slide_header(10, 'The Driver Experience', 'Score, risk, confidence, and plain-language reasons')

add_visual('App screenshots: Trip Results screen showing a circular score dial (85/100), risk level "Medium", confidence bar, and bullet-pointed reasons like "Frequent hard braking detected on asphalt roads." Below: Trip History list with route previews.')

add_script(
    'Let me show you what the driver actually sees. After completing a trip, the results screen displays a clear, circular safety score — '
    '85 out of 100 in this example. Next to it, the risk level: Low, Medium, or High. A confidence bar shows how certain the system is '
    'about its assessment. But the most important part is the reasons section. Instead of raw numbers or confusing model output, '
    'drivers see plain-language explanations: "Frequent hard braking detected on asphalt roads" or "Smooth acceleration maintained throughout trip." '
    'This transparency builds trust. Drivers understand what the system saw, and they know what to improve for next time. '
    'They can browse their trip history, replay routes on a map, and track their progress over time.'
)

add_tip('If possible, pull up a screenshot on your phone and gesture toward it as you describe each element.')

add_separator()

# ── Slide 11: Fleet Admin View ──────────────────────────────────────────────
add_slide_header(11, 'Fleet Management & Oversight', 'For operators who need visibility at scale')

add_visual('Admin dashboard showing: driver list with scores, flagged trip indicators, review queue with expandable trip details. Charts showing fleet-wide trends.')

add_script(
    'Drive Pulse is not just for individual drivers — it is built for fleet operators who need visibility at scale. '
    'The admin dashboard gives fleet managers a birds-eye view of their entire operation: all drivers listed with their scores, '
    'trips that need attention flagged for review, and expandable details for each trip. '
    'The review interface surfaces every trip\'s events, confidence data, and generated reasons in one clean workspace. '
    'This means a fleet manager in Addis Ababa can monitor drivers operating in Bahir Dar, Dire Dawa, or any other city, '
    'all from a single dashboard. And because admin access is role-based and restricted, you control who sees what.'
)

add_tip('Emphasize the "one dashboard, any city" angle — this is the scalability pitch that fleet operators care about most.')

add_separator()

# ── Slide 12: Languages & Accessibility ─────────────────────────────────────
add_slide_header(12, 'Built for Ethiopia', 'English · Amharic · Oromo — not an afterthought, a foundation')

add_visual('Three columns with language flags: English, Amharic (አማርኛ), Oromo (Afaan Oromoo). Each showing a UI element translated. Map of Ethiopia with major cities marked.')

add_script(
    'A critical design decision we made from day one: this product must work for Ethiopian drivers in their own languages. '
    'English, Amharic, and Oromo are all supported out of the box. The app interface, the trip results, the safety scores and reasons — '
    'everything is available in the driver\'s chosen language. '
    'Why does this matter? Because safety adoption depends on understanding. If a driver in rural Oromia cannot read the feedback '
    'because it is only in English, the system has failed. Language is not an afterthought for us — it is a foundation of the product. '
    'We believe this multilingual approach is essential for any safety technology deployed in Ethiopia and across Africa.'
)

add_tip('If you speak any of these languages, say a quick phrase in Amharic or Oromo to demonstrate. It connects with the judges personally.')

add_separator()

# ── Slide 13: Testimonials ──────────────────────────────────────────────────
add_slide_header(13, 'What Users Are Saying', 'Real feedback from our pilot participants')

add_visual('Three testimonial cards with photos/avatars and quotes.')

add_script(
    'But don\'t just take my word for it. Here is what our pilot participants have told us. One regular driver said: '
    '"I like seeing the trip score and plain reasons after each drive. It feels clear, not confusing, and it actually helps me notice my braking habits." '
    'A fleet reviewer told us: "The review side gives us one clean place to inspect flagged trips, confidence, and generated events. '
    'It is exactly what we needed." And an early product partner added: "What stands out is the combination of live trip capture '
    'and readable AI feedback. Even in the pilot, the product feels purposeful." '
    'This feedback validates our core design philosophy: transparency, usability, and real behavioral change.'
)

add_tip('Quotes are powerful. Read them with a slightly different tone, as if you are quoting someone directly. It adds authenticity.')

add_separator()

# ── Slide 14: Future Vision ─────────────────────────────────────────────────
add_slide_header(14, 'The Road Ahead', 'iOS, insurance partnerships, regional expansion')

add_visual('Three roadmap cards: (1) "iOS App" with Apple logo, (2) "Insurance Partnerships" with handshake icon, (3) "Regional Expansion" with East Africa map. Timeline arrow: Q3 2026 → Q1 2027 → Q4 2027.')

add_script(
    'So where do we go from here? Our roadmap has three major milestones. First, we are building the iOS version of the app to reach '
    'the significant number of iPhone users in Ethiopia and across Africa. Second, we are exploring partnerships with insurance companies '
    'to create usage-based insurance models — where your driving score directly translates to premium discounts. '
    'This would be a game-changer for safe drivers in Ethiopia. And third, we are planning regional expansion across East Africa, '
    'starting with Kenya and Uganda, where the same road safety challenges exist. '
    'But the vision goes beyond insurance. We see Drive Pulse becoming a platform — integrated with fleet management systems, '
    'public transportation authorities, and eventually, a standard feature in every smartphone sold in emerging markets.'
)

add_tip('The insurance partnership angle is particularly compelling for judges who care about sustainability and business models. Emphasize it.')

add_separator()

# ── Slide 15: Thank You & Closing ──────────────────────────────────────────
add_slide_header(15, 'Thank You', '"Technology cannot replace responsibility. But it can help prevent tragedy."')

add_visual('Closing slide with animated ambient glow. "Technology cannot replace responsibility. But it can help prevent tragedy." "DrivePulse — Building smarter, safer journeys through AI." Contact info at bottom.')

add_script(
    'I want to close with this thought: technology cannot replace responsibility. No app, no algorithm, no AI model can force a driver '
    'to make the right choice behind the wheel. But what technology can do — what Drive Pulse does — is give drivers the information, '
    'the feedback, and the awareness they need to make safer choices. It turns every trip into a learning opportunity. '
    'It gives fleet operators the visibility they need to protect their teams. And it does all of this using hardware that is already '
    'in people\'s pockets, in languages they actually speak. '
    'My name is [Your Name], and on behalf of Team VisionZero from Addis Ababa, Ethiopia — thank you for your time and your attention. '
    'I am happy to answer any questions you may have.'
)

add_tip('Slow down for the closing statement. Pause after "responsibility." Make eye contact with each judge as you say "thank you." End confidently with "I am happy to answer any questions."')

# ── Appendix: Quick Reference ───────────────────────────────────────────────
doc.add_page_break()
doc.add_heading('Quick Reference Card', level=1)

p = doc.add_paragraph()
run = p.add_run('Total presentation time: ')
run.bold = True
run = p.add_run('10–12 minutes (+ 5 minutes Q&A)')

doc.add_paragraph()
doc.add_paragraph('Key statistics to remember:', style='Heading 3')
stats = [
    '1.16 million global road deaths per year — one every 27 seconds',
    'Ethiopia: 37 per 100k fatality rate (double global average)',
    'Phone distraction: 22% of crashes; 6.1× risk when texting',
    '88% smartphone penetration in Africa by 2030',
    '50+ pilot trips, 144K+ sensor samples, 6 drivers, 6 event types',
    'Zero extra hardware — just a phone',
    '3 languages: English, Amharic, Oromo',
    'Live infrastructure on Render + Supabase PostgreSQL',
]
for s in stats:
    doc.add_paragraph(s, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Contact information:', style='Heading 3')
p = doc.add_paragraph('Website: drivepulse.onrender.com')
p = doc.add_paragraph('Email: sofoniaskidus@gmail.com')
p = doc.add_paragraph('Phone: +251 911 422 570')

# ── Save ────────────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'Drive_Pulse_Presentation_Script.docx')
doc.save(output_path)
print(f"Document saved to: {output_path}")
